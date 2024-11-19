from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
import os
from typing import Optional

app = FastAPI(debug=True)

# Configurar CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5000"],  # Orígenes permitidos
    allow_credentials=True,
    allow_methods=["*"],  # Métodos permitidos
    allow_headers=["*"],  # Encabezados permitidos
)

class EmployeeData(BaseModel):
    nombre: Optional[str] 
    apellidos: Optional[str] 
    cedula: Optional[str] 
    telefono: Optional[str] 
    cargo: Optional[str]
    fechaingreso: Optional[str]
    tipocontrato: Optional[str]
    direccion: Optional[str]
    salario: Optional[str]
    selector: Optional[str]
    
def modificar_contrato(nombre, apellidos, telefono, cedula, direccion, cargo, fechaingreso, tipocontrato, salario, output_path="contrato_modificado.docx"):
    plantilla_path = "contrato_plantilla.docx"
    doc = Document(plantilla_path)

    for paragraph in doc.paragraphs:
        if "{{NOMBRE}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{NOMBRE}}", nombre)
        if "{{APELLIDO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{APELLIDO}}", apellidos)
        if "{{TELEFONO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{TELEFONO}}}", telefono)
        if "{{CEDULA}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{CEDULA}}", cedula)
        if "{{DIRECCION}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{DIRECCION}}", direccion)    
        if "{{CARGO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{CARGO}}", cargo) 
        if "{{FECHAINGRESO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{FECHAINGRESO}}", fechaingreso)
        if "{{TIPOCONTRATO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{TIPOCONTRATO}}", tipocontrato)
        if "{{SALARIO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{SALARIO}}", salario)
    doc.save(output_path)
    print(f"Documento generado: {output_path}")

def modificar_afiliacion(nombre, apellidos, telefono, cedula, direccion, cargo, fechaingreso, tipocontrato, salario, output_path="contrato_modificado.docx"):
    plantilla_path = "contrato_plantilla.docx"
    doc = Document(plantilla_path)

    for paragraph in doc.paragraphs:
        if "{{NOMBRE}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{NOMBRE}}", nombre)
        if "{{APELLIDO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{APELLIDO}}", apellidos)
        if "{{TELEFONO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{TELEFONO}}}", telefono)
        if "{{CEDULA}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{CEDULA}}", cedula)
        if "{{DIRECCION}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{DIRECCION}}", direccion)    
        if "{{CARGO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{CARGO}}", cargo) 
        if "{{FECHAINGRESO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{FECHAINGRESO}}", fechaingreso)
        if "{{TIPOCONTRATO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{TIPOCONTRATO}}", tipocontrato)
        if "{{SALARIO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{SALARIO}}", salario)
    doc.save(output_path)
    print(f"Documento generado: {output_path}")

def modificar_certificado(nombre, apellidos, telefono, cedula, direccion, cargo, fechaingreso, tipocontrato, salario, output_path="contrato_modificado.docx"):
    plantilla_path = "contrato_plantilla.docx"
    doc = Document(plantilla_path)

    for paragraph in doc.paragraphs:
        if "{{NOMBRE}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{NOMBRE}}", nombre)
        if "{{APELLIDO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{APELLIDO}}", apellidos)
        if "{{TELEFONO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{TELEFONO}}}", telefono)
        if "{{CEDULA}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{CEDULA}}", cedula)
        if "{{DIRECCION}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{DIRECCION}}", direccion)    
        if "{{CARGO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{CARGO}}", cargo) 
        if "{{FECHAINGRESO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{FECHAINGRESO}}", fechaingreso)
        if "{{TIPOCONTRATO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{TIPOCONTRATO}}", tipocontrato)
        if "{{SALARIO}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{SALARIO}}", salario)
    doc.save(output_path)
    print(f"Documento generado: {output_path}")

@app.post("/documentos")
async def generar_contrato(data: EmployeeData):
    print(f"Datos recibidos: {data.model_dump()}")  # Depuración
    try:
        output_path = "contrato_modificado.docx"
        selector=data.selector
        if selector=='1':
            modificar_contrato(
                nombre=data.nombre,
                apellidos=data.apellidos,
                telefono=data.telefono,
                cedula=data.cedula,
                direccion=data.direccion,
                cargo=data.cargo,
                fechaingreso=data.fechaingreso,
                tipocontrato=data.tipocontrato,
                salario=data.salario,
                output_path=output_path
            )
        if selector=='2':
            modificar_afiliacion(
                nombre=data.nombre,
                apellidos=data.apellidos,
                telefono=data.telefono,
                cedula=data.cedula,
                direccion=data.direccion,
                cargo=data.cargo,
                fechaingreso=data.fechaingreso,
                tipocontrato=data.tipocontrato,
                salario=data.salario,
                output_path=output_path
            )
        if selector=='3':
            modificar_certificado(
                nombre=data.nombre,
                apellidos=data.apellidos,
                telefono=data.telefono,
                cedula=data.cedula,
                direccion=data.direccion,
                cargo=data.cargo,
                fechaingreso=data.fechaingreso,
                tipocontrato=data.tipocontrato,
                salario=data.salario,
                output_path=output_path
            )
        print("Contrato generado con éxito.")  # Registro de éxito
        if not os.path.exists(output_path):
            raise HTTPException(status_code=500, detail="Error al generar el archivo.")
        
        
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="contrato_modificado.docx",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


